import streamlit as st
import pandas as pd
import os
from docx import Document
import re
from collections import defaultdict
import tempfile
import io

# Use session state to avoid reprocessing the template on every interaction
if 'template_variables' not in st.session_state:
    st.session_state.template_variables = None
if 'template_doc' not in st.session_state:
    st.session_state.template_doc = None
if 'var_paragraphs_map' not in st.session_state:
    st.session_state.var_paragraphs_map = None
if 'csv_data' not in st.session_state:
    st.session_state.csv_data = None

def extract_variables_from_template(doc):
    """Extract all variables in {variable} format from the document more efficiently"""
    variables = set()
    pattern = r'\{([^{}]+)\}'
    var_paragraphs = defaultdict(list)
    
    # Process paragraphs and collect variables
    for i, paragraph in enumerate(doc.paragraphs):
        matches = re.findall(pattern, paragraph.text)
        for match in matches:
            variables.add(match)
            var_paragraphs[match].append(i)
    
    # Also check tables if they exist
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    matches = re.findall(pattern, paragraph.text)
                    variables.update(matches)
    
    return list(variables), var_paragraphs

def convert_to_csv(file):
    """Convert various spreadsheet formats to CSV data"""
    try:
        # Check file extension
        file_extension = os.path.splitext(file.name)[1].lower()
        
        # Read the file based on the extension
        if file_extension in ['.xlsx', '.xls', '.xlsm', '.xlsb']:
            # Excel files
            df = pd.read_excel(file)
        elif file_extension in ['.ods']:
            # OpenDocument Spreadsheet
            df = pd.read_excel(file, engine='odf')
        elif file_extension in ['.csv']:
            # CSV files
            df = pd.read_csv(file)
        elif file_extension in ['.tsv', '.txt']:
            # TSV or text files
            df = pd.read_csv(file, sep='\t')
        else:
            # Try to read as Excel by default
            df = pd.read_excel(file)
        
        # Convert to CSV
        csv_buffer = io.StringIO()
        df.to_csv(csv_buffer, index=False)
        
        return df, csv_buffer.getvalue()
    except Exception as e:
        st.error(f"Error converting file: {str(e)}")
        return None, None

def generate_documents(df, template_doc, column_mapping, filename_column, output_dir, var_paragraphs_map):
    """Generate individual documents based on Excel data with improved performance"""
    # Create output directory if it doesn't exist
    os.makedirs(output_dir, exist_ok=True)
    
    # Cache variable patterns for faster replacement
    var_patterns = {var: re.compile(r'\{' + re.escape(var) + r'\}') for var in column_mapping.keys()}
    
    # Process each row in the dataframe
    doc_count = 0
    for index, row in df.iterrows():
        # Create a new document from the template
        doc = Document()
        
        # Pre-calculate all replacements
        replacements = {var: str(row[column]) for var, column in column_mapping.items()}
        
        # Process paragraphs in the template
        for i, template_para in enumerate(template_doc.paragraphs):
            new_paragraph = doc.add_paragraph()
            new_paragraph.style = template_para.style
            
            # Only perform replacements if needed
            text = template_para.text
            
            # Check if this paragraph needs processing based on our map
            variables_to_replace = [var for var in column_mapping.keys() if i in var_paragraphs_map.get(var, [])]
            
            if variables_to_replace:
                for var in variables_to_replace:
                    pattern = var_patterns[var]
                    replacement_value = replacements[var]
                    text = pattern.sub(replacement_value, text)
            
            new_paragraph.text = text
        
        # Get filename from the selected column
        filename = str(row[filename_column]).replace(" ", "_")
        filename = re.sub(r'[^\w\-_\.]', '', filename)  # Remove invalid characters
        
        # Save the document
        output_path = os.path.join(output_dir, f"{filename}.docx")
        doc.save(output_path)
        doc_count += 1
    
    return doc_count

def main():
    st.title("Document Generator")
    
    # File uploads
    st.header("Step 1: Upload Files")
    
    template_file = st.file_uploader("Upload Word Template (with {variables})", type=["docx"], key="template_uploader")
    
    # Accept multiple spreadsheet formats
    spreadsheet_file = st.file_uploader(
        "Upload Spreadsheet Data (Excel, CSV, TSV, ODS, etc.)", 
        type=["xlsx", "xls", "csv", "tsv", "ods", "xlsm", "xlsb", "txt"],
        key="spreadsheet_uploader"
    )
    
    if template_file is not None and spreadsheet_file is not None:
        # Convert spreadsheet to CSV in the backend
        with st.spinner("Processing spreadsheet..."):
            df, csv_data = convert_to_csv(spreadsheet_file)
            
            if df is None:
                st.error("Failed to process the spreadsheet file. Please check the format.")
                st.stop()
            
            # Store CSV data in session state
            st.session_state.csv_data = csv_data
        
        # Show a preview of the data
        st.subheader("Data Preview")
        st.dataframe(df.head())
        
        # Process template only when it changes, using session state
        if template_file != st.session_state.get('last_template_file'):
            with st.spinner("Processing template..."):
                template_doc = Document(template_file)
                variables, var_paragraphs_map = extract_variables_from_template(template_doc)
                
                # Store in session state
                st.session_state.template_doc = template_doc
                st.session_state.template_variables = variables
                st.session_state.var_paragraphs_map = var_paragraphs_map
                st.session_state.last_template_file = template_file
        
        variables = st.session_state.template_variables
        template_doc = st.session_state.template_doc
        var_paragraphs_map = st.session_state.var_paragraphs_map
        
        if variables:
            st.header("Step 2: Map Variables to Excel Columns")
            
            # Create mapping UI
            column_mapping = {}
            available_columns = df.columns.tolist()
            
            # Create a more efficient UI for mapping with default selection when names match
            col1, col2 = st.columns(2)
            for i, var in enumerate(variables):
                # Determine which column to use for better UI layout
                current_col = col1 if i % 2 == 0 else col2
                
                default_index = 0
                # Try to find matching column name
                for j, col in enumerate(available_columns):
                    if var.lower() == col.lower():
                        default_index = j
                        break
                
                column_mapping[var] = current_col.selectbox(
                    f"Map '{var}' to column:", 
                    options=available_columns,
                    index=default_index,
                    key=f"mapping_{var}"
                )
            
            # Select naming column
            filename_column = st.selectbox(
                "Select column to use for output filenames:",
                options=available_columns,
                key="filename_column"
            )
            
            # Output directory
            output_dir = st.text_input("Output directory path:", "output_documents", key="output_dir")
            
            # Generate button
            if st.button("Generate Documents", key="generate_button"):
                with st.spinner("Generating documents..."):
                    num_docs = generate_documents(df, template_doc, column_mapping, filename_column, output_dir, var_paragraphs_map)
                st.success(f"{num_docs} documents generated successfully in '{output_dir}'!")
                
                # Option to download CSV data
                st.download_button(
                    label="Download Processed Data as CSV",
                    data=st.session_state.csv_data,
                    file_name="processed_data.csv",
                    mime="text/csv",
                )
        else:
            st.warning("No variables found in the template. Make sure to use {variable_name} format.")

if __name__ == "__main__":
    main()